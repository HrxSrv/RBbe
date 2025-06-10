from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import re
import unicodedata
from loguru import logger
from fastapi import HTTPException, status

# PDF processing imports
try:
    import PyPDF2
    import pdfplumber
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    logger.warning("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")

# DOC/DOCX processing imports
try:
    from docx import Document
    DOC_PROCESSING_AVAILABLE = True
except ImportError:
    DOC_PROCESSING_AVAILABLE = False
    logger.warning("python-docx not available. Install python-docx for DOC/DOCX support.")

class TextExtractionResult:
    """
    Structured result from text extraction operations.
    """
    def __init__(self, 
                 text: str, 
                 method: str, 
                 confidence: float, 
                 metadata: Dict[str, Any],
                 needs_vlm_processing: bool = False):
        self.text = text
        self.method = method  # "direct_pdf", "ocr_pdf", "docx", "plain_text"
        self.confidence = confidence  # 0.0 to 1.0
        self.metadata = metadata
        self.needs_vlm_processing = needs_vlm_processing
        self.extraction_timestamp = datetime.utcnow()
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "method": self.method,
            "confidence": self.confidence,
            "metadata": self.metadata,
            "needs_vlm_processing": self.needs_vlm_processing,
            "extraction_timestamp": self.extraction_timestamp.isoformat(),
            "text_length": len(self.text),
            "word_count": len(self.text.split()) if self.text else 0
        }

class TextExtractionService:
    """
    Service for extracting and preprocessing text from various document formats.
    Designed to work with Gemini VLM integration for complex documents.
    """
    
    # Text quality thresholds
    MIN_TEXT_LENGTH = 100  # Minimum characters for quality text
    MIN_WORD_COUNT = 20    # Minimum words for quality text
    MIN_CONFIDENCE_THRESHOLD = 0.7  # Below this, recommend VLM processing
    
    # Common resume section patterns for quality assessment
    RESUME_SECTION_PATTERNS = [
        r'\b(experience|education|skills|summary|objective|employment|work\s+history)\b',
        r'\b(projects|achievements|certifications|awards|publications)\b',
        r'\b(contact|email|phone|address|linkedin)\b'
    ]
    
    @classmethod
    async def extract_text(cls, file_path: str) -> TextExtractionResult:
        """
        Main entry point for text extraction from any supported file type.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            TextExtractionResult with extracted text and metadata
            
        Raises:
            HTTPException: If extraction fails or file type is unsupported
        """
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"File not found: {file_path}"
                )
            
            file_extension = file_path_obj.suffix.lower()
            
            # Route to appropriate extraction method
            if file_extension == ".pdf":
                return await cls._extract_from_pdf(file_path)
            elif file_extension in [".doc", ".docx"]:
                return await cls._extract_from_doc(file_path)
            elif file_extension == ".txt":
                return await cls._extract_from_text(file_path)
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file type: {file_extension}"
                )
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Text extraction failed. Please try again."
            )
    
    @classmethod
    async def _extract_from_pdf(cls, file_path: str) -> TextExtractionResult:
        """
        Extract text from PDF files using multiple strategies.
        """
        if not PDF_PROCESSING_AVAILABLE:
            logger.warning("PDF processing not available, recommending VLM processing")
            return TextExtractionResult(
                text="",
                method="pdf_unavailable",
                confidence=0.0,
                metadata={"error": "PDF processing libraries not available"},
                needs_vlm_processing=True
            )
        
        # Strategy 1: Try PyPDF2 first (fastest)
        pypdf_result = await cls._extract_pdf_with_pypdf2(file_path)
        if pypdf_result.confidence > cls.MIN_CONFIDENCE_THRESHOLD:
            return pypdf_result
        
        # Strategy 2: Try pdfplumber (more accurate for complex layouts)
        pdfplumber_result = await cls._extract_pdf_with_pdfplumber(file_path)
        if pdfplumber_result.confidence > pypdf_result.confidence:
            return pdfplumber_result
        
        # Strategy 3: If both methods produce low-quality text, recommend VLM
        best_result = pdfplumber_result if pdfplumber_result.confidence > pypdf_result.confidence else pypdf_result
        
        if best_result.confidence < cls.MIN_CONFIDENCE_THRESHOLD:
            best_result.needs_vlm_processing = True
            logger.info(f"PDF text extraction confidence low ({best_result.confidence:.2f}), recommending VLM processing")
        
        return best_result
    
    @classmethod
    async def _extract_pdf_with_pypdf2(cls, file_path: str) -> TextExtractionResult:
        """
        Extract text from PDF using PyPDF2 (fast but basic).
        """
        try:
            text = ""
            metadata = {"method": "pypdf2", "pages_processed": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["total_pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            metadata["pages_processed"] += 1
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
            
            # Clean and assess text quality
            cleaned_text = cls._clean_text(text)
            confidence = cls._assess_text_quality(cleaned_text)
            
            metadata.update({
                "raw_text_length": len(text),
                "cleaned_text_length": len(cleaned_text),
                "extraction_success_rate": metadata["pages_processed"] / metadata["total_pages"] if metadata["total_pages"] > 0 else 0
            })
            
            return TextExtractionResult(
                text=cleaned_text,
                method="direct_pdf_pypdf2",
                confidence=confidence,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            return TextExtractionResult(
                text="",
                method="direct_pdf_pypdf2",
                confidence=0.0,
                metadata={"error": str(e)}
            )
    
    @classmethod
    async def _extract_pdf_with_pdfplumber(cls, file_path: str) -> TextExtractionResult:
        """
        Extract text from PDF using pdfplumber (more accurate for complex layouts).
        """
        try:
            text = ""
            metadata = {"method": "pdfplumber", "pages_processed": 0}
            
            with pdfplumber.open(file_path) as pdf:
                metadata["total_pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            metadata["pages_processed"] += 1
                        
                        # Extract tables if present
                        tables = page.extract_tables()
                        if tables:
                            metadata["tables_found"] = metadata.get("tables_found", 0) + len(tables)
                            for table in tables:
                                # Convert table to text
                                table_text = cls._table_to_text(table)
                                text += table_text + "\n"
                    
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
            
            # Clean and assess text quality
            cleaned_text = cls._clean_text(text)
            confidence = cls._assess_text_quality(cleaned_text)
            
            metadata.update({
                "raw_text_length": len(text),
                "cleaned_text_length": len(cleaned_text),
                "extraction_success_rate": metadata["pages_processed"] / metadata["total_pages"] if metadata["total_pages"] > 0 else 0
            })
            
            return TextExtractionResult(
                text=cleaned_text,
                method="direct_pdf_pdfplumber",
                confidence=confidence,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            return TextExtractionResult(
                text="",
                method="direct_pdf_pdfplumber",
                confidence=0.0,
                metadata={"error": str(e)}
            )
    
    @classmethod
    async def _extract_from_doc(cls, file_path: str) -> TextExtractionResult:
        """
        Extract text from DOC/DOCX files using python-docx.
        """
        if not DOC_PROCESSING_AVAILABLE:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOC/DOCX processing not available. Please install python-docx."
            )
        
        try:
            text = ""
            metadata = {"method": "python-docx", "paragraphs_processed": 0, "tables_found": 0}
            
            doc = Document(file_path)
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    metadata["paragraphs_processed"] += 1
            
            # Extract table text
            for table in doc.tables:
                metadata["tables_found"] += 1
                table_text = ""
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text += " | ".join(row_text) + "\n"
                text += table_text + "\n"
            
            # Clean and assess text quality
            cleaned_text = cls._clean_text(text)
            confidence = cls._assess_text_quality(cleaned_text)
            
            metadata.update({
                "raw_text_length": len(text),
                "cleaned_text_length": len(cleaned_text)
            })
            
            return TextExtractionResult(
                text=cleaned_text,
                method="docx_extraction",
                confidence=confidence,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"DOC/DOCX extraction failed: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to extract text from document: {str(e)}"
            )
    
    @classmethod
    async def _extract_from_text(cls, file_path: str) -> TextExtractionResult:
        """
        Extract text from plain text files.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            cleaned_text = cls._clean_text(text)
            confidence = cls._assess_text_quality(cleaned_text)
            
            metadata = {
                "method": "plain_text",
                "raw_text_length": len(text),
                "cleaned_text_length": len(cleaned_text),
                "encoding": "utf-8"
            }
            
            return TextExtractionResult(
                text=cleaned_text,
                method="plain_text",
                confidence=confidence,
                metadata=metadata
            )
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    cleaned_text = cls._clean_text(text)
                    confidence = cls._assess_text_quality(cleaned_text)
                    
                    metadata = {
                        "method": "plain_text",
                        "raw_text_length": len(text),
                        "cleaned_text_length": len(cleaned_text),
                        "encoding": encoding
                    }
                    
                    return TextExtractionResult(
                        text=cleaned_text,
                        method="plain_text",
                        confidence=confidence,
                        metadata=metadata
                    )
                    
                except UnicodeDecodeError:
                    continue
            
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unable to decode text file. Unsupported encoding."
            )
        
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to extract text from file: {str(e)}"
            )
    
    @classmethod
    def _clean_text(cls, text: str) -> str:
        """
        Clean and normalize extracted text.
        """
        if not text:
            return ""
        
        # Normalize unicode characters
        text = unicodedata.normalize('NFKD', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Clean up common PDF artifacts
        text = re.sub(r'[\uf000-\uf8ff]', '', text)  # Remove private use area characters
        text = re.sub(r'[^\x00-\x7F]+', lambda m: m.group(0) if m.group(0).isprintable() else ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip and ensure we don't have empty result
        text = text.strip()
        
        return text
    
    @classmethod
    def _assess_text_quality(cls, text: str) -> float:
        """
        Assess the quality of extracted text to determine confidence.
        Returns value between 0.0 and 1.0.
        """
        if not text or len(text) < cls.MIN_TEXT_LENGTH:
            return 0.0
        
        word_count = len(text.split())
        if word_count < cls.MIN_WORD_COUNT:
            return 0.2
        
        confidence = 0.3  # Base confidence
        
        # Check for resume-like content
        resume_pattern_matches = 0
        for pattern in cls.RESUME_SECTION_PATTERNS:
            if re.search(pattern, text, re.IGNORECASE):
                resume_pattern_matches += 1
        
        # Boost confidence based on resume patterns
        pattern_score = min(resume_pattern_matches / len(cls.RESUME_SECTION_PATTERNS), 1.0)
        confidence += pattern_score * 0.4
        
        # Check text length and structure
        if len(text) > 500:
            confidence += 0.1
        if word_count > 100:
            confidence += 0.1
        
        # Check for email/phone patterns (common in resumes)
        if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):
            confidence += 0.05
        if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):
            confidence += 0.05
        
        # Penalize if too many special characters (might indicate OCR issues)
        special_char_ratio = len(re.findall(r'[^\w\s]', text)) / len(text) if text else 0
        if special_char_ratio > 0.1:
            confidence -= 0.1
        
        return min(max(confidence, 0.0), 1.0)
    
    @classmethod
    def _table_to_text(cls, table: List[List[str]]) -> str:
        """
        Convert extracted table data to readable text.
        """
        if not table:
            return ""
        
        text_lines = []
        for row in table:
            if row and any(cell and cell.strip() for cell in row):
                # Filter out empty cells and join with pipe separator
                clean_cells = [cell.strip() for cell in row if cell and cell.strip()]
                if clean_cells:
                    text_lines.append(" | ".join(clean_cells))
        
        return "\n".join(text_lines)
    
    @classmethod
    async def batch_extract_text(cls, file_paths: List[str]) -> Dict[str, TextExtractionResult]:
        """
        Extract text from multiple files in batch.
        """
        results = {}
        
        for file_path in file_paths:
            try:
                result = await cls.extract_text(file_path)
                results[file_path] = result
            except Exception as e:
                logger.error(f"Batch extraction failed for {file_path}: {e}")
                results[file_path] = TextExtractionResult(
                    text="",
                    method="batch_extraction_failed",
                    confidence=0.0,
                    metadata={"error": str(e)},
                    needs_vlm_processing=True
                )
        
        return results
    
    @classmethod
    def get_extraction_summary(cls, results: Dict[str, TextExtractionResult]) -> Dict[str, Any]:
        """
        Generate summary statistics for batch extraction results.
        """
        if not results:
            return {}
        
        total_files = len(results)
        successful_extractions = sum(1 for r in results.values() if r.confidence > 0)
        vlm_recommended = sum(1 for r in results.values() if r.needs_vlm_processing)
        
        avg_confidence = sum(r.confidence for r in results.values()) / total_files
        total_text_length = sum(len(r.text) for r in results.values())
        
        methods_used = {}
        for result in results.values():
            methods_used[result.method] = methods_used.get(result.method, 0) + 1
        
        return {
            "total_files": total_files,
            "successful_extractions": successful_extractions,
            "success_rate": successful_extractions / total_files,
            "vlm_processing_recommended": vlm_recommended,
            "vlm_recommendation_rate": vlm_recommended / total_files,
            "average_confidence": avg_confidence,
            "total_text_extracted": total_text_length,
            "methods_used": methods_used,
            "timestamp": datetime.utcnow().isoformat()
        } 